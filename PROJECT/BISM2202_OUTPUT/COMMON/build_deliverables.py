#!/usr/bin/env python3
"""Build BISM2202 reports, guides, DAX, QA files, and status documents.

Re-run this script after adding real Power BI screenshots to either screenshots/
folder. Existing screenshots are inserted into the corresponding DOCX; missing
screenshots remain explicit placeholders. No PBIX or fake images are created.
"""

from __future__ import annotations

import argparse
import json
import os
import platform
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
COMMON = ROOT / "COMMON"
SOURCE_DOCX_NAME = "Data Visualization Using Microsoft Power BI Assessment Task Instructions 2026 Semester 2_BISM2202.docx"
SOURCE_XLSX_NAME = "Assignment1_BISM2202_pizza_sell_data.xlsx"


def discover_source(filename: str, environment_name: str) -> Path:
    """Resolve an input without depending on a user-specific macOS path."""
    configured = os.environ.get(environment_name)
    script = Path(__file__).resolve()
    roots = [Path.cwd(), ROOT, *ROOT.parents, Path.home() / "Downloads"]
    candidates = ([Path(configured).expanduser()] if configured else []) + [
        root / relative
        for root in roots
        for relative in (filename, Path("INPUTS") / filename, Path("input") / filename)
    ]
    for candidate in candidates:
        if candidate.is_file():
            return candidate.resolve()
    return Path(filename)


SOURCE_DOCX = discover_source(SOURCE_DOCX_NAME, "BISM2202_SOURCE_DOCX")
SOURCE_XLSX = discover_source(SOURCE_XLSX_NAME, "BISM2202_SOURCE_XLSX")
RESULTS = json.loads((COMMON / "analysis_results.json").read_text(encoding="utf-8"))
TOTAL = RESULTS["metadata"]["total_orders"]

QUESTIONS = {
    1: "What is the top 20 locations based on count of order number?",
    2: "What is the average Delivery Duration (min) by Pizza Size?",
    3: "What is the share of orders by Pizza Size?",
    4: "How many orders were placed with each Payment Method?",
    5: "What is the share of orders by Traffic Level?",
    6: "What is the average Toppings Count for weekend vs. weekday orders?",
    7: "Which locations have the highest average Delivery Duration (min)?",
    8: "How many orders were placed during peak hours vs. non-peak?",
    9: "How does order volume trend by Month-Year?",
    10: "What is the average Toppings Count by Pizza Size?",
    11: "Which 2 restaurants have the lowest average Delay (min) i.e., the fastest, most reliable performers?",
    12: "Which 5 Pizza Types have the highest average Delivery Duration (min)?",
    13: "How does order volume vary by Order Hour across a typical day?",
    14: "What is the average delay (min) by Traffic Level?",
    15: "What proportion of orders fall into each Pizza Complexity category?",
    16: "Build a table of average Delivery Duration and Delay by Restaurant Name and use the Fx (Format by field value) button to visually highlight the fastest and slowest restaurants.",
    17: "How does order volume vary by Order Hour, alongside average Delay (min)?",
    18: "How does average Topping Density vary by Pizza Type, alongside order volume?",
    19: "What is the percentage breakdown of orders by Traffic Level within each Payment methods, filterable by order time?",
    20: "Design three different charts or tables that use conditional formatting (or highlight them by using different colors) in a single dashboard. Explain why each one matters, what it represents, how they should be read together, and what insights emerge from viewing them as a whole.",
}

POINTS = {**{i: 0.5 for i in range(1, 9)}, **{i: 1 for i in range(9, 17)}, **{i: 2 for i in range(17, 21)}}

VISUALS = {
    "A": {
        1: ("Clustered bar chart", "Y-axis: Location; X-axis: [Order Count]", "Top N = 20 by [Order Count]; descending"),
        2: ("Clustered column chart", "X-axis: Pizza Size; Y-axis: [Avg Delivery Duration]", "Data labels on; 1 decimal"),
        3: ("Donut chart", "Legend: Pizza Size; Values: [Order Count]", "Detail labels: category + percent of total"),
        4: ("Clustered column chart", "X-axis: Payment Method; Y-axis: [Order Count]", "Descending by [Order Count]"),
        5: ("Donut chart", "Legend: Traffic Level; Values: [Order Count]", "Category + percentage labels"),
        6: ("Clustered column chart", "X-axis: Weekend Label; Y-axis: [Avg Toppings Count]", "Data labels on; 2 decimals"),
        7: ("Clustered bar chart", "Y-axis: Location; X-axis: [Avg Delivery Duration]", "Top N = 10 by [Avg Delivery Duration]; descending; retain ties"),
        8: ("Donut chart", "Legend: Peak Hour Label; Values: [Order Count]", "Category + percentage labels"),
        9: ("Line chart", "X-axis: Order Month-Year; Y-axis: [Order Count]", "Order Month-Year ascending from 2024-01 to 2026-07"),
        10: ("Clustered column chart", "X-axis: Pizza Size; Y-axis: [Avg Toppings Count]", "Data labels on; 2 decimals"),
        11: ("Clustered bar chart", "Y-axis: Restaurant Name; X-axis: [Avg Delay]", "Visual filter [Delay Rank Asc] <= 2; ascending"),
        12: ("Clustered bar chart", "Y-axis: Pizza Type; X-axis: [Avg Delivery Duration]", "Top N = 5 by [Avg Delivery Duration]; descending"),
        13: ("Line chart", "X-axis: Order Hour; Y-axis: [Order Count]", "Order Hour ascending; data markers on"),
        14: ("Clustered column chart", "X-axis: Traffic Level; Y-axis: [Avg Delay]", "Low/Medium/High risk colors; data labels on"),
        15: ("Donut chart", "Legend: Pizza Complexity; Values: [Order Count]", "Treat numeric field as category; percentage labels"),
        16: ("Table", "Columns: Restaurant Name, [Avg Delivery Duration], [Avg Delay]", "Background color Fx > Field value using two color measures"),
        17: ("Line and clustered column chart", "X-axis: Order Hour; Column Y-axis: [Order Count]; Line Y-axis: [Avg Delay]", "Secondary Y-axis on; hour ascending"),
        18: ("Line and clustered column chart", "X-axis: Pizza Type; Column Y-axis: [Order Count]; Line Y-axis: [Avg Topping Density]", "Secondary Y-axis on; descending by [Order Count]"),
        19: ("100% stacked column chart + slicer", "X-axis: Payment Method; Legend: Traffic Level; Y-axis: [Order Count]; slicer: Order Time", "Slicer style Between; test interaction"),
        20: ("Dashboard: table + column + combo", "Restaurant performance; Traffic delay; Hourly volume and delay", "Three visuals; conditional formatting/intentional colors"),
    },
    "B": {
        1: ("Clustered column chart", "X-axis: Location; Y-axis: [Order Count]", "Top N = 20; descending; labels readable"),
        2: ("Clustered bar chart", "Y-axis: Pizza Size; X-axis: [Avg Delivery Duration]", "Data labels on; 1 decimal"),
        3: ("Treemap", "Group: Pizza Size; Values: [Order Count]", "Category + percentage labels"),
        4: ("Clustered bar chart", "Y-axis: Payment Method; X-axis: [Order Count]", "Descending by [Order Count]"),
        5: ("Treemap", "Group: Traffic Level; Values: [Order Count]", "Category + percentage labels"),
        6: ("Clustered bar chart", "Y-axis: Weekend Label; X-axis: [Avg Toppings Count]", "Data labels on; 2 decimals"),
        7: ("Clustered column chart", "X-axis: Location; Y-axis: [Avg Delivery Duration]", "Top N = 10; descending; retain ties"),
        8: ("Clustered column chart", "X-axis: Peak Hour Label; Y-axis: [Order Count]", "Data labels on"),
        9: ("Clustered column chart", "X-axis: Order Month-Year; Y-axis: [Order Count]", "Order Month-Year ascending from 2024-01 to 2026-07"),
        10: ("Clustered bar chart", "Y-axis: Pizza Size; X-axis: [Avg Toppings Count]", "Data labels on; 2 decimals"),
        11: ("Clustered column chart", "X-axis: Restaurant Name; Y-axis: [Avg Delay]", "Visual filter [Delay Rank Asc] <= 2; ascending"),
        12: ("Clustered column chart", "X-axis: Pizza Type; Y-axis: [Avg Delivery Duration]", "Top N = 5; descending"),
        13: ("Area chart", "X-axis: Order Hour; Y-axis: [Order Count]", "Order Hour ascending; markers optional"),
        14: ("Clustered bar chart", "Y-axis: Traffic Level; X-axis: [Avg Delay]", "Cool-to-warm ordered colors; labels on"),
        15: ("Treemap", "Group: Pizza Complexity; Values: [Order Count]", "Category + percentage labels"),
        16: ("Matrix", "Rows: Restaurant Name; Values: [Avg Delivery Duration], [Avg Delay]", "Fx > Field value using Version B color measures"),
        17: ("Line and stacked column chart", "X-axis: Order Hour; Column Y-axis: [Order Count]; Line Y-axis: [Avg Delay]", "Secondary Y-axis on; hour ascending"),
        18: ("Line and clustered column chart", "X-axis: Pizza Type; Column Y-axis: [Order Count]; Line Y-axis: [Avg Topping Density]", "Secondary axis on; stronger density emphasis"),
        19: ("100% stacked bar chart + slicer", "Y-axis: Payment Method; Legend: Traffic Level; X-axis: [Order Count]; slicer: Order Time", "Slicer style Between; test interaction"),
        20: ("Dashboard: bar + column + 100% stacked bar", "Top locations; Pizza type delivery time; Traffic mix by payment", "Order Time slicer; three coordinated visuals"),
    },
}

TITLES = {
    "A": ["Top 20 Locations by Order Volume", "Average Delivery Duration by Pizza Size", "Share of Orders by Pizza Size", "Orders by Payment Method", "Share of Orders by Traffic Level", "Average Toppings Count: Weekend vs Weekday", "Locations with Highest Average Delivery Duration", "Peak vs Non-Peak Order Volume", "Order Volume Trend by Month-Year", "Average Toppings Count by Pizza Size", "Two Restaurants with Lowest Average Delay", "Top 5 Pizza Types by Average Delivery Duration", "Hourly Order Volume", "Average Delay by Traffic Level", "Order Share by Pizza Complexity", "Restaurant Delivery Performance", "Hourly Order Volume and Average Delay", "Pizza Type Topping Density and Order Volume", "Traffic Mix within Payment Method", "Pizza Delivery Operations Dashboard"],
    "B": ["Top 20 Order Locations", "Delivery Time across Pizza Sizes", "Pizza Size Order Distribution", "Payment Method Order Ranking", "Traffic-Level Order Mix", "Topping Levels on Weekends and Weekdays", "Slowest Delivery Locations", "Orders in Peak and Non-Peak Periods", "Order Volume by Month-Year", "Toppings across Pizza Sizes", "Most Reliable Restaurants by Average Delay", "Pizza Types with Longest Delivery Times", "Orders across a Typical Day", "Traffic Conditions and Average Delay", "Pizza Complexity Mix", "Restaurant Speed and Delay Matrix", "When Order Pressure and Delay Rise", "Pizza Type Volume and Topping Density", "Traffic Composition inside Payment Methods", "Pizza Order and Delivery Performance Overview"],
}


def frame(key: str) -> pd.DataFrame:
    return pd.DataFrame(RESULTS["questions"][key]["result"])


def narratives() -> dict[str, dict[int, str]]:
    a = {
        1: "Atlanta, GA ranked first with 78 orders, followed by Milwaukee, WI with 71 and Louisville, KY with 69. The horizontal ranking makes the difference between the leading locations easy to compare.",
        2: "Large pizzas had the longest average delivery duration at 34.17 minutes, while Small pizzas had the shortest at 21.86 minutes. XL pizzas were close to Large pizzas at 33.08 minutes.",
        3: "Medium pizzas accounted for the largest share, with 429 orders (42.73%). Small pizzas had the smallest share, with 132 orders (13.15%).",
        4: "Card was the most common payment method with 276 orders, only five more than UPI at 271. Hut Points and Domino's Cash were much less common, at 24 and 23 orders respectively.",
        5: "Medium traffic represented the largest share of orders at 39.64% (398 orders). High traffic accounted for 32.67%, while Low traffic represented 27.69%.",
        6: "Weekday orders averaged 3.40 toppings, compared with 3.27 on weekends. The difference was small, at about 0.13 toppings per order.",
        7: "Fort Wayne, IN and Newark, NJ had the highest average delivery duration at 50.00 minutes. Laredo, Lexington, Minneapolis, and Orlando followed at 45.00 minutes, so several locations shared the next-highest result.",
        8: "Peak hours contained 949 orders (94.52%), compared with only 55 non-peak orders (5.48%). Most observations in this dataset therefore fall in the peak-hour category.",
        9: "The chronological series contains 31 distinct Month-Year points from 2024-01 to 2026-07. Volume peaked at 86 orders in 2024-08 and remained high at 75 in 2024-09, before falling to 44 in 2024-10. The lowest complete point was 2024-05 with 6 orders; the final 2026-07 point has 7 orders and covers data only through 7 July.",
        10: "XL pizzas had the highest average toppings count at 4.99, followed by Large pizzas at 3.97. Small pizzas had the lowest average at 1.69 toppings.",
        11: "Little Caesars recorded the lowest average delay at 16.62 minutes. Domino's was second at 17.02 minutes, so these are the two lowest-delay restaurants under the assignment definition.",
        12: "Stuffed Crust had the highest average delivery duration at 39.52 minutes. Gluten-Free (32.44), Cheese Burst (32.26), Thai Chicken (31.67), and Sicilian (30.86) completed the top five.",
        13: "The largest hourly order volume occurred at 19:00 with 328 orders. The nearby 18:00 and 20:00 periods were also busy, with 312 and 306 orders, while the remaining observed hours were much quieter.",
        14: "Average delay increased across the traffic categories: 15.46 minutes for Low, 17.86 for Medium, and 19.16 for High traffic. The pattern shows an association between heavier traffic and longer delay in this dataset.",
        15: "Pizza Complexity 6 was the largest category, with 308 orders (30.68%). Complexity 12 represented 22.81% and Complexity 20 represented 20.12%, while categories 3 and 15 each contained only two orders.",
        16: "The conditional formatting identifies Papa John's as the fastest restaurant by average delivery duration at 28.19 minutes, while Domino's was slowest at 30.26 minutes. Little Caesars had the lowest average delay at 16.62 minutes, whereas Marco's Pizza had the highest at 18.44 minutes. The fastest delivery-duration restaurant is therefore not the same as the lowest-delay restaurant.",
        17: "Order volume peaked at 19:00 with 328 orders, but the highest average delay occurred at 20:00 at 19.97 minutes. The volume and delay peaks did not fully coincide. The 21:00 delay was also high at 19.75 minutes, but that hour contained only three orders and should be read cautiously.",
        18: "Non-Veg had the largest order volume at 216 and an average topping density of 0.760. Cheese Burst combined high volume (188 orders) with the highest average density, 0.845. The supplied density field equals toppings per kilometre, so it should not be interpreted as a pure pizza-size measure.",
        19: "Traffic composition differed by payment method. High traffic made up 44.57% of Card orders, Medium traffic was largest for Cash (48.02%) and UPI (47.60%), and Low traffic was largest for Wallet (44.71%). Hut Points was 95.83% High traffic, but it contained only 24 orders, so its percentage is based on a small group.",
        20: "The restaurant table identifies where delivery time and delay are highest or lowest, the traffic chart shows that average delay rises from Low to High traffic, and the hourly combo chart shows when operational pressure is greatest. Read together, they separate three useful views: restaurant performance, delivery conditions, and time-of-day demand.\n\nPapa John's had the shortest average delivery duration (28.19 minutes), while Marco's Pizza had the highest average delay (18.44 minutes). High traffic also had the highest average delay (19.16 minutes), and the order-volume peak at 19:00 was followed by the delay peak at 20:00. These patterns point to specific restaurants and periods for closer operational review without claiming that the visualized relationships prove causation.",
    }
    b = {
        1: "The compact ranking keeps Atlanta, GA as the leading location at 78 orders, ahead of Milwaukee, WI (71) and Louisville, KY (69). The alternative orientation emphasizes the size gap among the top demand locations.",
        2: "The bars place Large pizzas at the longest average delivery time, 34.17 minutes. Small pizzas are clearly lower at 21.86 minutes, while XL and Large are relatively close.",
        3: "The treemap gives the largest area to Medium pizzas, which generated 429 orders or 42.73% of the total. Small pizzas occupy the smallest area at 132 orders or 13.15%.",
        4: "The horizontal ranking shows Card (276) and UPI (271) as the two dominant payment methods. Wallet and Cash follow at 208 and 202, while the two restaurant-specific methods each have fewer than 25 orders.",
        5: "Medium traffic forms the largest block at 398 orders (39.64%). High traffic contributes 328 orders (32.67%) and Low traffic contributes 278 (27.69%).",
        6: "Weekday orders average 3.40 toppings and weekend orders average 3.27. The short bars are close, confirming that the observed difference is only about 0.13.",
        7: "Fort Wayne and Newark form the highest pair at 50.00 minutes. Four more locations share 45.00 minutes, so the column view should retain and clearly label tied results at the filter boundary.",
        8: "The columns show a strong imbalance: 949 Peak Hour orders versus 55 Non-Peak Hour orders. Peak periods represent 94.52% of all orders in the source.",
        9: "The 31 Month-Year columns run chronologically from 2024-01 to 2026-07. The largest point is 2024-08 with 86 orders, followed by 2024-09 with 75; the series then falls to 44 in 2024-10. The 2024-05 point is lowest at 6 orders, while 2026-07 contains 7 orders through 7 July only.",
        10: "XL stands out with an average of 4.99 toppings, while Small averages only 1.69. Large and Medium fall between them at 3.97 and 2.77.",
        11: "Little Caesars is the first low-delay performer at 16.62 minutes, followed by Domino's at 17.02. The two-column comparison uses the same bottom-two ranking as Version A.",
        12: "Stuffed Crust is separated from the other top-five pizza types at 39.52 minutes. The next four range from 30.86 to 32.44 minutes, led by Gluten-Free.",
        13: "The area is concentrated around 18:00–20:00. Order volume reaches 328 at 19:00, with 312 at 18:00 and 306 at 20:00; other observed hours contribute relatively few orders.",
        14: "The ordered bars rise from 15.46 minutes under Low traffic to 17.86 under Medium and 19.16 under High traffic. The color progression reinforces the observed association without treating it as causal proof.",
        15: "Complexity 6 occupies the largest treemap area with 30.68% of orders. Complexity 12 and 20 add 22.81% and 20.12%, so these three categories together dominate the mix.",
        16: "The matrix highlights different leaders for the two measures. Papa John's is fastest on average delivery duration at 28.19 minutes, while Little Caesars has the lowest average delay at 16.62 minutes. Domino's has the longest average duration (30.26), and Marco's Pizza has the highest average delay (18.44).",
        17: "The alternative combo styling shows that demand reaches its maximum at 19:00 (328 orders), while average delay reaches its maximum one hour later at 20:00 (19.97 minutes). This separation is easier to see when the delay line is emphasized on its own axis. The 21:00 value is based on only three orders.",
        18: "Cheese Burst has the highest average topping density at 0.845 while still recording 188 orders. Non-Veg has the greatest volume at 216 with density 0.760, whereas Gluten-Free is lower at 0.582. The supplied metric is toppings per kilometre, which limits a pizza-only interpretation.",
        19: "The 100% bars compare composition rather than payment-method size. Card is 44.57% High traffic; Cash and UPI are mainly Medium traffic at 48.02% and 47.60%; Wallet is mainly Low traffic at 44.71%. Small groups such as Hut Points (24 orders) and Domino's Cash (23) require cautious percentage comparisons.",
        20: "The dashboard begins with demand location, then compares pizza-type delivery performance, and finishes with traffic composition inside payment methods. Atlanta leads location demand with 78 orders, while Stuffed Crust has the longest pizza-type average delivery duration at 39.52 minutes.\n\nThe traffic-payment view adds operating context: Card is weighted toward High traffic, while Cash and UPI are weighted toward Medium traffic and Wallet toward Low traffic. Using the Order Time slicer lets the viewer test whether these patterns remain stable across different periods. Together, the three visuals connect where demand occurs, which pizza types take longer, and the conditions surrounding each payment mix.",
    }
    return {"A": a, "B": b}


ANSWERS = narratives()


def write_assignment_requirements() -> None:
    doc = Document(SOURCE_DOCX)
    dictionary = [(r.cells[0].text.strip(), r.cells[1].text.strip()) for r in doc.tables[0].rows[1:]]
    lines = ["# BISM2202 Assessment Requirements (extracted from instructor DOCX)", "",
             f"- Source: `{SOURCE_DOCX}`", "- Course: BISM2202", "- Assessment: Data Visualization Using Microsoft Power BI - Assessment Task",
             "- Total points: 20", "- Number of questions: 20", "- Scenario: pizza shop chain analysis; instructor text mentions operations in 2024–2025.", "",
             "## Submission requirements", "", "1. Submit the report to Turnitin.",
             "2. Submit both the report and the Power BI `.pbix` file to Blackboard.",
             "3. A submission without the `.pbix` attachment will not be marked.",
             "4. For every question, insert the relevant Power BI visualization screenshot and then write the observation/answer/explanation.", "",
             "## Questions and marks", ""]
    for i in range(1, 21):
        lines.append(f"- Q{i} ({POINTS[i]:g} pt): {QUESTIONS[i]}")
    lines += ["", "## Data dictionary", "", "| Column | Description |", "|---|---|"]
    lines.extend([f"| {c} | {d} |" for c, d in dictionary])
    lines += ["", "## Answer format", "",
              "The instructor document provides alternating fields for each question: first the visualization screenshot/image, then the written observation/answer/explanation. No word-count requirement appears in the instructor DOCX.", "",
              "## PBIX and screenshot requirements", "",
              "A real Power BI `.pbix` must be submitted with the report. The instructor asks for screenshots of the relevant data visualizations for all 20 questions. This workspace does not fabricate either artifact."]
    (COMMON / "ASSIGNMENT_REQUIREMENTS.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_dax() -> None:
    text = r'''# DAX Measures and Calculated Columns

Use the table name `PizzaOrders`. Before writing DAX, Power Query must replace the three occurrences of `Marco’s Pizza` with `Marco's Pizza` in `Restaurant Name`.

## Measures

```DAX
Order Count =
DISTINCTCOUNT(PizzaOrders[Order ID])

Avg Delivery Duration =
AVERAGE(PizzaOrders[Delivery Duration (min)])

Avg Toppings Count =
AVERAGE(PizzaOrders[Toppings Count])

Avg Delay =
AVERAGE(PizzaOrders[Delay (min)])

Avg Topping Density =
AVERAGE(PizzaOrders[Topping Density])

Order Share % =
DIVIDE(
    [Order Count],
    CALCULATE([Order Count], ALLSELECTED(PizzaOrders))
)

Traffic % Within Payment Method =
DIVIDE(
    [Order Count],
    CALCULATE(
        [Order Count],
        REMOVEFILTERS(PizzaOrders[Traffic Level])
    )
)

Delay Rank Asc =
RANKX(
    ALL(PizzaOrders[Restaurant Name]),
    [Avg Delay],
    ,
    ASC,
    DENSE
)
```

## Calculated columns

```DAX
Order Month-Year =
FORMAT(PizzaOrders[Order Time], "yyyy-MM")

Peak Hour Label =
IF(PizzaOrders[Is Peak Hour] = TRUE(), "Peak Hour", "Non-Peak Hour")

Weekend Label =
IF(PizzaOrders[Is Weekend] = TRUE(), "Weekend", "Weekday")
```

Because the label uses ISO `yyyy-MM`, sorting `Order Month-Year` ascending is also chronological. Do not use the source `Order Month` name alone, because that would combine the same month across different years.

## Version A field-value colors for Q16

```DAX
Delivery Duration Color A =
VAR CurrentValue = [Avg Delivery Duration]
VAR MinValue = MINX(ALL(PizzaOrders[Restaurant Name]), CALCULATE([Avg Delivery Duration]))
VAR MaxValue = MAXX(ALL(PizzaOrders[Restaurant Name]), CALCULATE([Avg Delivery Duration]))
RETURN SWITCH(TRUE(), CurrentValue = MinValue, "#C6EFCE", CurrentValue = MaxValue, "#FFC7CE", "#FFF2CC")

Delay Color A =
VAR CurrentValue = [Avg Delay]
VAR MinValue = MINX(ALL(PizzaOrders[Restaurant Name]), CALCULATE([Avg Delay]))
VAR MaxValue = MAXX(ALL(PizzaOrders[Restaurant Name]), CALCULATE([Avg Delay]))
RETURN SWITCH(TRUE(), CurrentValue = MinValue, "#C6EFCE", CurrentValue = MaxValue, "#FFC7CE", "#FFF2CC")
```

## Version B field-value colors for Q16

```DAX
Delivery Duration Color B =
VAR CurrentValue = [Avg Delivery Duration]
VAR MinValue = MINX(ALL(PizzaOrders[Restaurant Name]), CALCULATE([Avg Delivery Duration]))
VAR MaxValue = MAXX(ALL(PizzaOrders[Restaurant Name]), CALCULATE([Avg Delivery Duration]))
RETURN SWITCH(TRUE(), CurrentValue = MinValue, "#BFE3E0", CurrentValue = MaxValue, "#F4A261", "#E9EEF3")

Delay Color B =
VAR CurrentValue = [Avg Delay]
VAR MinValue = MINX(ALL(PizzaOrders[Restaurant Name]), CALCULATE([Avg Delay]))
VAR MaxValue = MAXX(ALL(PizzaOrders[Restaurant Name]), CALCULATE([Avg Delay]))
RETURN SWITCH(TRUE(), CurrentValue = MinValue, "#CDE8F6", CurrentValue = MaxValue, "#E76F51", "#EEF1F4")
```

## Expected global checks

- `[Order Count]` with no filter = 1,004.
- `[Avg Delivery Duration]` with no filter = 29.4920318725 minutes.
- `[Avg Delay]` with no filter = 17.6225498008 minutes.
- `Traffic % Within Payment Method` must sum to 100% inside every Payment Method and must respond to the Order Time slicer.
'''
    (COMMON / "DAX_MEASURES.md").write_text(text, encoding="utf-8")


def expected_check(i: int) -> str:
    checks = {
        1: "Atlanta, GA = 78; Milwaukee, WI = 71; Louisville, KY = 69.",
        2: "Large = 34.17; Medium = 27.53; Small = 21.86; XL = 33.08 minutes.",
        3: "Medium = 429 (42.73%); Small = 132 (13.15%); total = 1,004.",
        4: "Card 276; UPI 271; Wallet 208; Cash 202; Hut Points 24; Domino's Cash 23.",
        5: "Low 278 (27.69%); Medium 398 (39.64%); High 328 (32.67%).",
        6: "Weekday = 3.40; Weekend = 3.27.",
        7: "Fort Wayne and Newark = 50.00; four locations = 45.00; retain visible ties.",
        8: "Peak Hour = 949; Non-Peak Hour = 55.",
        9: "31 chronological points from 2024-01 to 2026-07; 2024-08 = 86; 2024-09 = 75; total = 1,004.",
        10: "Small 1.69; Medium 2.77; Large 3.97; XL 4.99.",
        11: "Little Caesars = 16.62; Domino's = 17.02.",
        12: "Stuffed Crust 39.52; Gluten-Free 32.44; Cheese Burst 32.26; Thai Chicken 31.67; Sicilian 30.86.",
        13: "18:00 = 312; 19:00 = 328; 20:00 = 306.",
        14: "Low 15.46; Medium 17.86; High 19.16.",
        15: "Complexity 6 = 308 (30.68%); 12 = 229 (22.81%); 20 = 202 (20.12%).",
        16: "Fastest duration Papa John's 28.19; slowest Domino's 30.26; lowest delay Little Caesars 16.62; highest delay Marco's Pizza 18.44.",
        17: "Order peak 19:00 = 328; delay peak 20:00 = 19.97.",
        18: "Non-Veg volume 216/density 0.760; Cheese Burst volume 188/density 0.845.",
        19: "Each Payment Method totals 100%; Card High = 44.57%; Cash Medium = 48.02%; UPI Medium = 47.60%; Wallet Low = 44.71%.",
        20: "Exactly three main visuals, differentiated colors or conditional formatting, and an integrated explanation based on validated earlier results.",
    }
    return checks[i]


def write_report_markdown(version: str) -> None:
    folder = ROOT / f"Version_{version}"
    lines = [f"# BISM2202 Power BI Assessment - Version {version}", "",
             "Student Name: __________", "", "Student ID: __________", ""]
    for i in range(1, 21):
        lines += [f"## Q{i}", "", QUESTIONS[i], "",
                  f"[Insert Power BI screenshot: screenshots/Q{i:02d}.png]", "",
                  "Observation:", "", ANSWERS[version][i], ""]
    (folder / "report_answers.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def set_cell_border_shading(paragraph, fill: str = "F3F6F9") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{side}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "8")
        tag.set(qn("w:color"), "C8D1DC")
        borders.append(tag)
    p_pr.append(borders)


def configure_doc(doc: Document, version: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, before, after in [("Heading 1", 16, 14, 8), ("Heading 2", 13, 12, 6), ("Heading 3", 12, 8, 4)]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(31, 78, 121)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    header = section.header.paragraphs[0]
    header.text = f"BISM2202 | Power BI Assessment | Version {version}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(100, 110, 120)
    footer = section.footer.paragraphs[0]
    footer.text = "Student Name: __________    |    Student ID: __________"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(100, 110, 120)


def build_docx(version: str) -> None:
    folder = ROOT / f"Version_{version}"
    doc = Document()
    configure_doc(doc, version)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DATA VISUALIZATION USING\nMICROSOFT POWER BI")
    r.font.name = "Calibri"
    r.font.size = Pt(25)
    r.font.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(36)
    r = p.add_run(f"BISM2202 Assessment Report - Version {version}")
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(70, 78, 88)
    for label in ["Student Name: __________", "Student ID: __________", "Total Questions: 20", "Total Points: 20"]:
        p = doc.add_paragraph(label)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph("This report presents the completed Power BI visualisation and evidence for each assessment question.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    for run in p.runs:
        run.font.italic = True
        run.font.color.rgb = RGBColor(90, 98, 108)

    for i in range(1, 21):
        doc.add_page_break()
        h = doc.add_heading(f"Q{i}  |  {POINTS[i]:g} point{'s' if POINTS[i] != 1 else ''}", level=1)
        h.paragraph_format.keep_with_next = True
        qp = doc.add_paragraph(QUESTIONS[i])
        qp.paragraph_format.space_after = Pt(10)
        qp.runs[0].font.bold = True
        image_path = folder / "screenshots" / f"Q{i:02d}.png"
        if image_path.exists() and image_path.stat().st_size > 0:
            pic_p = doc.add_paragraph()
            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_p.add_run().add_picture(str(image_path), width=Inches(6.55))
            cap = doc.add_paragraph(f"Figure Q{i}. Power BI visualization.")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = RGBColor(90, 98, 108)
        else:
            ph = doc.add_paragraph()
            ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ph.paragraph_format.space_before = Pt(10)
            ph.paragraph_format.space_after = Pt(18)
            ph.paragraph_format.line_spacing = 1.4
            ph.add_run(f"\n[Insert Q{i:02d} genuine Power BI screenshot here]\n\n").bold = True
            set_cell_border_shading(ph)
        label = doc.add_paragraph()
        label.paragraph_format.space_after = Pt(4)
        r = label.add_run("Observation / answer / explanation")
        r.bold = True
        r.font.color.rgb = RGBColor(31, 78, 121)
        for idx, paragraph in enumerate(ANSWERS[version][i].split("\n\n")):
            p = doc.add_paragraph(paragraph)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.widow_control = True
    out = folder / f"BISM2202_Report_{version}.docx"
    doc.save(out)


def common_guide_intro(version: str) -> list[str]:
    return [
        f"# Power BI 零基础操作指南 - Version {version}", "",
        "> 本指南以 Power BI Desktop 2026 界面为基准。新版可能启用 On-object interaction：如果右侧没有旧版 Visualizations pane，请选中 Visual 后直接在图上点 Build a visual / Add data；格式设置可双击 Visual 或右键 Format。两种界面使用相同字段槽位。", "",
        "## 0. 你需要准备", "",
        "1. 一台 Windows 10/11 x64 电脑或 Windows 虚拟机。macOS 没有原生 Power BI Desktop。",
        "2. 从 Microsoft Store 安装 **Microsoft Power BI Desktop**。打开后选择空白报告。",
        f"3. 把整个 `BISM2202_OUTPUT` 目录和原始 Excel `{SOURCE_XLSX.name}` 复制到 Windows 本地磁盘。不要改原始 Excel。",
        "4. 确认右侧能看到 Data、Visualizations/Build、Filters 三个区域。若看不到，打开 View（视图）选项卡，在 Panes 中勾选相应窗格。", "",
        "## 1. 导入 Excel 并执行唯一的数据清洗", "",
        "1. 在 Power BI Desktop 顶部功能区点 **Home > Get data > Excel workbook**。",
        f"2. 选择 `{SOURCE_XLSX.name}`，点 **Open**。",
        "3. Navigator 左侧勾选 `Sheet1`。右侧预览应显示 25 个字段。",
        "4. 不要直接点 Load；点 **Transform Data**，进入 Power Query Editor。Load 会直接载入，Transform Data 可先修正字段名和值。",
        "5. 左侧 Queries 对 `Sheet1` 右键 > Rename，输入 `PizzaOrders`。",
        "6. 选中 `Restaurant Name` 列，点 **Transform > Replace Values**。Value to Find 粘贴弯引号版本 `Marco’s Pizza`，Replace With 输入直引号版本 `Marco's Pizza`，点 OK。Applied Steps 应新增 Replaced Value。只会影响 3 行。",
        "7. 逐列检查类型图标：Order ID/Restaurant/Location/Pizza Size/Pizza Type/Traffic Level/Payment Method/Order Month/Payment Category 为 Text；Order Time/Delivery Time 为 Date/Time；Is Peak Hour/Is Weekend/Is Delayed 为 True/False；Toppings Count、Pizza Complexity、Traffic Impact、Order Hour 为 Whole Number；其余数值字段为 Decimal Number（Delivery Duration 也可为 Whole Number）。",
        "8. 点 **Home > Close & Apply**。返回报告画布后，Data pane 应出现 `PizzaOrders`。",
        "9. 在左侧 Table view（表格图标）检查总行数为 1,004；不要过滤或删除 2026 行。", "",
        "## 2. 创建 Measures 和 Columns", "",
        "1. 在右侧 Data pane 对 `PizzaOrders` 右键 > **New measure**；也可先选中表，再点 Home > Calculations > New measure。顶部会出现公式栏。",
        "2. 逐个复制 `COMMON/DAX_MEASURES.md` 中的 Measures。每粘贴一条按 Enter，确认表下出现计算器图标。",
        "3. 对 `PizzaOrders` 右键 > **New column**，逐个创建 `Order Month-Year`、`Peak Hour Label`、`Weekend Label`。`Order Month-Year = FORMAT(PizzaOrders[Order Time], \"yyyy-MM\")`。",
        "4. Q09 必须使用 `Order Month-Year` 并设为 Ascending；不要使用源字段 `Order Month`，否则不同年份的同名月份会被错误合并。",
        "5. 选中各平均值 Measure，在 Measure tools 把 Format 设为 Decimal number；Avg Delivery/Avg Delay 设 2 位，Avg Toppings/Avg Topping Density 设 2 或 3 位。Order Count 设 Whole number。", "",
        "## 3. 新增与重命名页面", "",
        "1. 报告底部点 `+` 新建页面。双击页面标签，依次创建并命名 Q01、Q02、…、Q19、Q20 Dashboard。",
        "2. 每个 Q01–Q19 页面只放该题核心 Visual；Q19 另加 Order Time slicer；Q20 放 3 个核心 Visual。",
        "3. 每做完一页立即 Ctrl+S。第一次保存时选择 Save As，Version A 保存为 `BISM2202_Assignment_A.pbix`，Version B 保存为 `BISM2202_Assignment_B.pbix`。", "",
        "## 4. 通用界面操作", "",
        "- 插入 Visual：先点画布空白处，再点 Visualizations pane 中的图表图标。新版 On-object 界面可点画布上的 Build a visual 再选择图表类型。",
        "- 拖字段：从 Data pane 展开 `PizzaOrders`，把列或 Measure 拖到 Build visual 的 X-axis、Y-axis、Legend、Values、Rows、Column Y-axis、Line Y-axis 等槽位。",
        "- 改汇总：字段槽位右边点下拉箭头，选 Count、Distinct count 或 Average。本指南优先拖已验证的 Measures，避免默认 Sum。",
        "- 改标题/标签：选中 Visual > Format visual（油漆刷）> General > Title；数据标签通常在 Visual > Data labels。",
        "- 排序：Visual 右上角 `...` > Sort axis/Sort by > 指定字段 > Ascending/Descending。",
        "- Visual filter：先选 Visual，再在 Filters pane 的 Filters on this visual 中展开字段卡。Top N 选择 Top、输入 N、把 Measure 拖到 By value，点 Apply filter。",
        "- 截图：确保只截完整 Visual、标题、轴/图例、标签；使用 Windows Snipping Tool，保存到对应 `screenshots/Qxx.png`。", "",
        "## 5. Q1–Q20 独立步骤", "",
    ]


def question_steps(version: str, i: int) -> list[str]:
    visual, slots, settings = VISUALS[version][i]
    lines = [f"### Q{i:02d} - {QUESTIONS[i]}", "",
             f"**目标 Visual：** {visual}  ", f"**字段槽位：** {slots}  ",
             f"**筛选/排序/格式：** {settings}  ", f"**标题：** `{TITLES[version][i-1]}`", "",
             f"1. 在底部单击页面 `{'Q20 Dashboard' if i == 20 else f'Q{i:02d}'}`，再单击画布空白处，确保没有选中别的 Visual。"]
    if i == 20:
        if version == "A":
            lines += [
                "2. 插入 Table：加入 Restaurant Name、[Avg Delivery Duration]、[Avg Delay]；按 Q16 的 Fx 方法给两个数值列设置 Field value 背景色。放在左侧，占页面约 40% 宽。",
                "3. 插入 Clustered column chart：X-axis 放 Traffic Level，Y-axis 放 [Avg Delay]；标题改为 `Average Delay by Traffic Level`，用 Low/Medium/High 三档颜色。放在右上。",
                "4. 插入 Line and clustered column chart：X-axis 放 Order Hour，Column Y-axis 放 [Order Count]，Line Y-axis 放 [Avg Delay]；开启 Secondary y-axis。放在右下。",
                "5. 插入 Text box 作为页面标题 `Pizza Delivery Operations Dashboard`。三个主 Visual 必须同时完整可读。",
            ]
        else:
            lines += [
                "2. 插入 Clustered bar chart：Y-axis 放 Location，X-axis 放 [Order Count]；Filters on this visual 设 Top 10 by [Order Count]，降序。放在左上。",
                "3. 插入 Clustered column chart：X-axis 放 Pizza Type，Y-axis 放 [Avg Delivery Duration]；设 Top 5 by [Avg Delivery Duration]，降序。放在左下。",
                "4. 插入 100% stacked bar chart：Y-axis 放 Payment Method，Legend 放 Traffic Level，X-axis 放 [Order Count]。在右侧插入 Slicer，字段放 Order Time，Format > Slicer settings > Options > Style 选 Between。",
                "5. 插入 Text box 作为页面标题 `Pizza Order and Delivery Performance Overview`。调整布局，使三个主 Visual 和 slicer 都不重叠。",
            ]
        lines += [
            "6. 单击不同图形部分，确认交叉高亮有意义；若某个 Visual 变成空白，检查字段槽位和 Visual-level filters。",
            f"7. 核对 Python 基准：{expected_check(i)}",
            f"8. 保存 PBIX，用 Snipping Tool 截取完整 dashboard，保存为 `screenshots/Q{i:02d}.png`。", ""]
        return lines
    lines += [f"2. 在 Visualizations/Build visual 中选择 **{visual}**。若图标名称不确定，把鼠标停在图标上等待工具提示。",
              f"3. 从 Data pane 展开 `PizzaOrders`，按以下位置逐个拖字段：**{slots}**。Measure 带计算器图标，普通列没有。"]
    if i in (1, 7, 12):
        n = 20 if i == 1 else (10 if i == 7 else 5)
        measure = "[Order Count]" if i == 1 else "[Avg Delivery Duration]"
        lines += [f"4. 保持 Visual 选中，在 Filters pane > Filters on this visual 中展开类别字段，Filter type 选 Top N；Show items 选 Top；输入 {n}；把 {measure} 拖到 By value；点 Apply filter。",
                  f"5. 点 Visual 右上 `...` > Sort by {measure} > Descending。若 Q7 因并列显示超过 10 个位置，这是正常的 tie 行为，两个版本必须保留相同并列项。"]
    elif i == 11:
        lines += ["4. 把 Measure `[Delay Rank Asc]` 拖到 Filters on this visual，设置 `is less than or equal to 2`，点 Apply filter。",
                  "5. Visual 右上 `...` > Sort by [Avg Delay] > Ascending，只应显示 Little Caesars 与 Domino's。"]
    elif i == 19:
        lines += ["4. 单击画布空白处，插入 Slicer visual；把 `Order Time`（不要用 Date hierarchy）拖入 Field。",
                  "5. 选中 slicer，Format visual > Visual > Slicer settings > Options > Style 选 Between。拖动左右端点，确认 100% stacked chart 会变化，再恢复完整日期范围。"]
    elif i == 16:
        color1 = f"Delivery Duration Color {version}"
        color2 = f"Delay Color {version}"
        lines += [f"4. 在 Build visual 的 Values 区，点 `[Avg Delivery Duration]` 右侧下拉箭头 > Conditional formatting > Background color（新版也可到 Format visual > Cell elements 找 Fx）。",
                  f"5. Conditional formatting 对话框中 Format style 选 Field value；What field should we base this on/Based on field 选 `[{color1}]`；Summarization 选 First；点 OK。",
                  f"6. 对 `[Avg Delay]` 重复：Background color > Fx > Format style = Field value > Based on field = `[{color2}]` > OK。确认不是手工逐格上色。"]
    else:
        lines += [f"4. 设置 Visual：{settings}。若字段默认显示 Sum，点字段槽位右侧下拉箭头，改用指定 Measure 或 Average/Distinct count。",
                  f"5. 点 Visual 右上 `...`，按题意完成排序；然后在 Format visual 中把 Data labels 开启，并按需要设置 1–2 位小数。"]
    if i in (3, 5, 8, 15):
        lines.append("6. 在 Format visual > Detail labels/Data labels 中开启 Category 和 Percent of total；若选项名称不同，展开 Labels 的内容选项，确保截图同时能看出类别和百分比。")
    elif i in (17, 18):
        lines.append("6. 在 Format visual 展开 Secondary y-axis，设为 On，并开启左右轴标题；Order Count 使用整数轴，平均值使用 1–2 位小数。")
    elif i == 9:
        lines.append("6. X-axis 必须使用 `Order Month-Year`（显示为 yyyy-MM），不是源字段 `Order Month`；Visual 右上 `...` > Sort axis > Order Month-Year > Ascending。应出现 31 个时间点，从 2024-01 到 2026-07。")
    elif i == 13:
        lines.append("6. 确认 X-axis 使用 Whole Number `Order Hour`，不是把它当日期或文本；排序必须是 12、13、14、17、18、19、20、21。")
    else:
        lines.append("6. 检查轴标题、图例和数据标签没有遮挡；必要时拉宽 Visual，不要缩小到看不清字段名。")
    lines += [f"7. 在 Format visual > General > Title 开启标题并输入 `{TITLES[version][i-1]}`。",
              f"8. 核对 Python 基准：{expected_check(i)}",
              f"9. Ctrl+S 保存 PBIX；用 Windows Snipping Tool 截完整 Visual，保存为 `screenshots/Q{i:02d}.png`。", ""]
    return lines


def write_guides(version: str) -> None:
    lines = common_guide_intro(version)
    for i in range(1, 21):
        lines.extend(question_steps(version, i))
    lines += ["## 6. 最终保存、截图和报告插图", "",
              f"1. 确认 20 个页面名称完整；Q19 slicer 可互动；Q16 两个字段都使用 Fx > Field value；Q20 有三张主要 Visual。",
              f"2. File > Save As，把 PBIX 保存到 `Version_{version}/BISM2202_Assignment_{version}.pbix`。关闭后重新打开一次，确认数据和 Visual 没有错误。",
              "3. 检查 `screenshots/Q01.png` 到 `Q20.png` 都是真实 Power BI 截图。不要使用 Python/Excel 图代替。",
              "4. 把项目复制回本机后，在项目根目录运行 bundled Python：`python BISM2202_OUTPUT/COMMON/build_deliverables.py`。脚本会自动把已存在的截图插入 DOCX。",
              f"5. 打开 `BISM2202_Report_{version}.docx`，逐题核对截图数字、标题和英文 observation 完全一致。", "",
              "## 官方界面参考", "",
              "- https://learn.microsoft.com/en-us/power-bi/create-reports/desktop-excel-stunning-report",
              "- https://learn.microsoft.com/en-us/power-bi/create-reports/power-bi-on-object-interaction",
              "- https://learn.microsoft.com/en-us/power-bi/visuals/power-bi-visualization-slicer-visual",
              "- https://learn.microsoft.com/en-us/power-bi/create-reports/desktop-conditional-table-formatting",
              "- https://learn.microsoft.com/en-us/power-bi/visuals/power-bi-visualization-combo-chart", ""]
    path = ROOT / f"Version_{version}" / f"POWERBI_ZERO_BEGINNER_GUIDE_{version}.md"
    path.write_text("\n".join(lines), encoding="utf-8")


def write_checklists(version: str) -> None:
    folder = ROOT / f"Version_{version}"
    shot = [f"# Screenshot Checklist - Version {version}", "",
            "所有方框必须在 Windows Power BI Desktop 中逐项确认。当前截图状态：`PENDING_WINDOWS_POWERBI`。", ""]
    for i in range(1, 21):
        shot += [f"## Q{i:02d}", "", "- [ ] Title visible", "- [ ] Labels readable",
                 "- [ ] Correct filter", "- [ ] Correct sort", "- [ ] No field error",
                 "- [ ] No visual warning", "- [ ] Screenshot complete and from real Power BI Desktop", ""]
    (folder / f"SCREENSHOT_CHECKLIST_{version}.md").write_text("\n".join(shot), encoding="utf-8")

    qa = [f"# QA Checklist - Version {version}", "",
          "Python基准已完成；Power BI、截图和 PBIX 项必须在 Windows 上验证。", ""]
    for i in range(1, 21):
        visual = VISUALS[version][i][0]
        qa += [f"## Q{i:02d}", "", f"- [x] Question requirement mapped: {QUESTIONS[i]}",
               f"- [x] Correct planned visual: {visual}", "- [x] Correct field/aggregation/filter/sort documented",
               f"- [x] Python verification target: {expected_check(i)}",
               "- [x] Report text uses validated numbers", "- [ ] Power BI visual matches Python",
               "- [ ] Screenshot present and readable", "- [ ] Final report screenshot and text match", ""]
    qa += ["## Special gates", "",
           "- [ ] Q01 Top 20", "- [ ] Q07 highest average duration and ties handled consistently",
           "- [ ] Q11 bottom 2 average delay", "- [ ] Q12 top 5 average duration",
           "- [ ] Q16 Fx > Format by field value used for both value columns",
           "- [ ] Q17 Order Count + Avg Delay with secondary axis",
           "- [ ] Q18 Order Count + Avg Topping Density with secondary axis",
           "- [ ] Q19 100% stacked + Order Time Between slicer + tested interaction",
           "- [ ] Q20 exactly three main visuals + differentiated colors/conditional formatting + integrated explanation",
           "- [ ] PBIX reopens without errors and contains all 20 pages", ""]
    (folder / f"QA_CHECKLIST_{version}.md").write_text("\n".join(qa), encoding="utf-8")


def write_cross_validation() -> None:
    metrics = [
        "Total Orders", "Top 20 Locations", "Average Delivery Duration by Pizza Size", "Pizza Size share",
        "Payment Method counts", "Traffic Level share", "Weekend / Weekday Avg Toppings",
        "Location Avg Delivery ranking", "Peak / Non-Peak counts", "Month-Year order totals",
        "Avg Toppings by Pizza Size", "Bottom 2 restaurants by Avg Delay",
        "Top 5 Pizza Types by Avg Delivery Duration", "Hourly Order Count", "Avg Delay by Traffic Level",
        "Pizza Complexity share", "Restaurant Avg Delivery Duration", "Restaurant Avg Delay",
        "Order Hour Avg Delay", "Pizza Type Avg Topping Density", "Payment Method x Traffic Level percentages",
    ]
    lines = ["# Cross-Version Validation", "", "**Overall baseline result: PASS**", "",
             "Version A and Version B reports, guides, DAX, and QA targets are generated from the same `analysis_results.json`. No version-specific data filters or alternative calculations are used. This PASS covers the analytical baseline; final PBIX/screenshot reconciliation remains pending on Windows.", "",
             "| Metric | Version A baseline | Version B baseline | Result |", "|---|---|---|---|"]
    for metric in metrics:
        lines.append(f"| {metric} | Shared Python result | Shared Python result | PASS |")
    lines += ["", "## Fixed shared rules", "",
              "- Source rows: 1,004; DISTINCTCOUNT(Order ID): 1,004.",
              "- Restaurant typography normalization: `Marco’s Pizza` -> `Marco's Pizza` for 3 rows, in memory/Power Query only.",
              "- Q9 uses 31 distinct Month-Year points from 2024-01 to 2026-07 and sorts them chronologically.",
              "- Q19 percentages use Payment Method as denominator and must respond to the Order Time slicer.",
              "- If either PBIX differs from `analysis_results.xlsx`, mark `FAIL_DATA_INCONSISTENCY`, correct the visual/filter/aggregation, and repeat the check."]
    (COMMON / "CROSS_VERSION_VALIDATION.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_environment_and_status() -> None:
    env = f'''# Environment Status

- OS: macOS 26.6.2 (Build 25G83)
- CPU architecture: arm64
- Hardware: Apple M4 Max, 16 cores, 48 GB RAM
- System Python: 3.9.6
- Bundled workspace Python: available and used for analysis/document generation
- Microsoft Excel app: not installed
- Microsoft Word app: not installed
- LibreOffice app: not installed
- Bundled headless LibreOffice: available for DOCX render QA
- Microsoft Power BI Desktop: not installed; no native macOS Power BI Desktop exists
- Windows: No
- Windows ARM: No
- Virtualization support: hardware support present; no Parallels/VMware/VirtualBox/UTM guest process detected

## Can complete reliably here

- Read the instructor DOCX and source XLSX without modifying either file.
- Validate data quality and calculate Q1–Q20 support in Python.
- Generate JSON/XLSX baselines, DAX, two design specifications, report text, DOCX drafts, beginner guides, and QA files.
- Render the DOCX files with bundled headless LibreOffice and inspect the page images.

## Cannot complete reliably here

- Create genuine Power BI report pages in Power BI Desktop.
- Capture genuine Power BI screenshots.
- Create or reopen genuine `.pbix` files.

These three items remain `PENDING_WINDOWS_POWERBI`. No placeholder `.pbix` or fake visualization image is created.
'''
    (ROOT / "ENVIRONMENT_STATUS.md").write_text(env, encoding="utf-8")
    ai = '''# AI Writing Review Guide

The report text is a short, data-based draft. Before submitting:

1. Read every answer aloud and confirm you understand what each number means.
2. Compare every number with the final Power BI visual and `COMMON/analysis_results.xlsx`.
3. Adjust a few expressions to match your normal English, but do not change the statistical meaning.
4. Do not add causes that the dataset cannot prove. Use “is associated with” rather than “caused”.
5. Confirm you can explain the aggregation, filter, and visual choice to the instructor.
6. Remove every screenshot placeholder and insert only genuine Power BI Desktop screenshots.
7. Do not use hidden characters, “humanizers”, deliberate mistakes, or promises about an AI-detection percentage.
8. Follow your university’s academic-integrity and permitted-AI-use rules, including any disclosure requirement.
'''
    (ROOT / "AI_WRITING_REVIEW_GUIDE.md").write_text(ai, encoding="utf-8")
    status = f'''# FINAL_STATUS

## Environment

- OS: **COMPLETE** - macOS 26.6.2, Apple M4 Max / arm64.
- Power BI Desktop: **BLOCKED** - not available natively on macOS and not installed.
- Can build PBIX here: **BLOCKED** - genuine PBIX creation requires Windows Power BI Desktop.

## COMMON

- Data read: **COMPLETE** - `{SOURCE_XLSX}` / Sheet1 / 1,004 rows / 25 columns.
- Instructor requirements read: **COMPLETE** - `{SOURCE_DOCX}`; 20 questions, 20 points, screenshot plus explanation for each question, report + PBIX submission.
- Data quality check: **COMPLETE** - quality findings and one typography normalization documented.
- Python validation: **COMPLETE** - reproducible script, JSON, and XLSX created.
- Q1-Q20 calculated: **COMPLETE** - Q1-Q19 numeric outputs plus Q20 validated support metrics.
- Cross-version validation: **COMPLETE** - PASS for the shared analytical baseline.

## Version_A

- Design: **COMPLETE** - traditional business analytics specification for Q01-Q20.
- Report answers: **COMPLETE** - data-backed English draft.
- DOCX: **COMPLETE** - screenshot placeholders are intentional until genuine images exist.
- Power BI pages: **PENDING** - requires Windows Power BI Desktop.
- Screenshots: **PENDING** - `PENDING_WINDOWS_POWERBI`; no fake screenshots created.
- PBIX: **PENDING** - `PENDING_WINDOWS_POWERBI`; no fake PBIX created.
- QA: **PENDING** - checklist is complete, but final visual/PBIX checks require Windows.

## Version_B

- Design: **COMPLETE** - distinct alternative visual system using the same data.
- Report answers: **COMPLETE** - data-backed English draft with visual-specific emphasis.
- DOCX: **COMPLETE** - screenshot placeholders are intentional until genuine images exist.
- Power BI pages: **PENDING** - requires Windows Power BI Desktop.
- Screenshots: **PENDING** - `PENDING_WINDOWS_POWERBI`; no fake screenshots created.
- PBIX: **PENDING** - `PENDING_WINDOWS_POWERBI`; no fake PBIX created.
- QA: **PENDING** - checklist is complete, but final visual/PBIX checks require Windows.

## MANUAL_STEPS_REMAINING

1. Use a Windows 10/11 x64 computer or a Windows virtual machine.
2. Install Microsoft Power BI Desktop.
3. Copy the source XLSX and the entire `BISM2202_OUTPUT` directory to Windows.
4. Choose Version A or Version B and open its `POWERBI_ZERO_BEGINNER_GUIDE_*.md`.
5. Import Sheet1, rename it `PizzaOrders`, and apply the documented 3-row Restaurant Name replacement in Power Query.
6. Create the DAX measures/columns and Q01-Q20 pages exactly as documented.
7. Save genuine screenshots to `screenshots/Q01.png` through `Q20.png`.
8. Save and reopen the genuine PBIX at the required filename.
9. Copy the project back and rerun `COMMON/build_deliverables.py` to insert screenshots into the DOCX.
10. Complete every QA/screenshot checkbox, review the Word report, and submit the report plus PBIX according to Turnitin/Blackboard instructions.
'''
    (ROOT / "FINAL_STATUS.md").write_text(status, encoding="utf-8")


def main() -> None:
    global SOURCE_DOCX, SOURCE_XLSX
    parser = argparse.ArgumentParser(description="Build the BISM2202 Word reports and support files.")
    parser.add_argument("--source-docx", type=Path, default=SOURCE_DOCX)
    parser.add_argument("--source-xlsx", type=Path, default=SOURCE_XLSX)
    args = parser.parse_args()

    SOURCE_DOCX = args.source_docx.resolve()
    SOURCE_XLSX = args.source_xlsx.resolve()
    if not SOURCE_DOCX.is_file():
        raise FileNotFoundError(f"Instructor DOCX not found: {SOURCE_DOCX}")
    if not SOURCE_XLSX.is_file():
        raise FileNotFoundError(f"Source XLSX not found: {SOURCE_XLSX}")

    write_assignment_requirements()
    write_dax()
    write_cross_validation()
    write_environment_and_status()
    for version in ("A", "B"):
        write_report_markdown(version)
        write_guides(version)
        write_checklists(version)
        build_docx(version)
    print("Built requirements, DAX, guides, reports, QA checklists, DOCX files, and status documents.")


if __name__ == "__main__":
    main()
